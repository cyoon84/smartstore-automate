#!/usr/bin/env python3
"""
스마트스토어 출고 엑셀에서 구매자별 주문 품목 리스트(.docx)를 생성.

사용법:
    python3 generate.py <스마트스토어_경로> <비밀번호>

영문 상품명 우선순위:
  1) ~/smartstore-project/.cowork-skills/order-list/preferred-names.json (상품번호||옵션정보 키)
  2) ~/smartstore-project/templates/product-mapping.xlsx 의 상품명(영문)
  3) 한글 상품명 (폴백)

표준출력:
  DONE:<output_path>:<buyer_count>:<line_count>
  MISSING_NAMES:상품번호1|상품번호2...   (이름을 못 찾은 상품이 있을 때)
"""
import io
import json
import os
import sys
from collections import OrderedDict
from datetime import datetime

import msoffcrypto
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


HERE = os.path.dirname(os.path.abspath(__file__))
PREFERRED_PATH = os.path.join(os.path.dirname(HERE), "preferred-names.json")
PRODUCT_MAPPING = os.path.expanduser("~/smartstore-project/templates/product-mapping.xlsx")
OUTPUT_DIR = os.path.expanduser("~/smartstore-project/output")


def load_preferred():
    if not os.path.exists(PREFERRED_PATH):
        return {}
    with open(PREFERRED_PATH, encoding="utf-8") as f:
        return json.load(f)


def load_eng_map():
    if not os.path.exists(PRODUCT_MAPPING):
        return {}
    mp = pd.read_excel(PRODUCT_MAPPING, dtype={"상품번호": str})
    mp["상품번호"] = mp["상품번호"].astype(str).str.split(".").str[0]
    mp = mp.drop_duplicates(subset="상품번호", keep="first")
    return dict(zip(mp["상품번호"], mp["상품명(영문)"]))


def read_smartstore(path, password):
    decrypted = io.BytesIO()
    with open(os.path.expanduser(path), "rb") as f:
        if password:
            of = msoffcrypto.OfficeFile(f)
            of.load_key(password=password)
            of.decrypt(decrypted)
            decrypted.seek(0)
        else:
            decrypted = io.BytesIO(f.read())
    return pd.read_excel(decrypted, engine="openpyxl", header=1)


def resolve_name(pn, option, kor, preferred, eng_map):
    key = f"{pn}||{option}"
    if key in preferred:
        return preferred[key]["preferred_eng"], False
    if pn in eng_map:
        v = eng_map[pn]
        if pd.notna(v) and str(v).strip():
            return str(v).strip(), False
    return kor, True  # missing eng


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color="CCCCCC", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcb = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcb.append(b)
    tc_pr.append(tcb)


def main():
    if len(sys.argv) < 2:
        print("Usage: generate.py <smartstore_path> [password]", file=sys.stderr)
        sys.exit(1)

    smartstore_path = sys.argv[1]
    password = sys.argv[2] if len(sys.argv) > 2 else ""

    preferred = load_preferred()
    eng_map = load_eng_map()

    ss = read_smartstore(smartstore_path, password)
    ss["_pn"] = ss["상품번호"].astype(str).str.split(".").str[0]
    ss["_opt"] = ss["옵션정보"].apply(lambda v: "" if pd.isna(v) else str(v).strip())

    missing_set = []

    # buyer → list of (eng, qty, kor, pn, opt)
    by_buyer = OrderedDict()
    # (buyer, pn, opt) 단위로 수량 합산해 한 줄로 표시
    agg = OrderedDict()
    recipients = {}  # buyer → set of recipients

    for _, r in ss.iterrows():
        buyer = str(r["구매자명"]).strip()
        pn = r["_pn"]
        opt = r["_opt"]
        kor = str(r["상품명"])
        qty = int(r["수량"]) if not pd.isna(r["수량"]) else 1
        eng, missing = resolve_name(pn, opt, kor, preferred, eng_map)
        if missing:
            missing_set.append(pn)

        key = (buyer, pn, opt)
        if key in agg:
            agg[key]["qty"] += qty
        else:
            agg[key] = {"buyer": buyer, "eng": eng, "qty": qty, "kor": kor, "pn": pn, "opt": opt}

        recipients.setdefault(buyer, set()).add(str(r["수취인명"]).strip())

    # buyer 별로 항목 리스트 구성, 가나다 정렬
    for entry in agg.values():
        by_buyer.setdefault(entry["buyer"], []).append(entry)
    sorted_buyers = sorted(by_buyer.keys())

    today = datetime.now().strftime("%Y%m%d")
    title_date = datetime.now().strftime("%Y/%-m/%-d")  # 2026/5/7 형식
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    out_path = os.path.join(OUTPUT_DIR, f"핀치마트_{today}.docx")

    # ---- docx 생성 ----
    doc = Document()
    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    # 페이지 여백 살짝 줄임
    for section in doc.sections:
        section.top_margin = Cm(1.9)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)

    # 제목
    h = doc.add_paragraph()
    run = h.add_run(f"핀치마트 {title_date} order list")
    run.bold = True
    run.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub_run = sub.add_run(
        f"생성일: {datetime.now().strftime('%Y-%m-%d')}    구매자 {len(sorted_buyers)}명 / 주문 라인 {sum(len(v) for v in by_buyer.values())}건"
    )
    sub_run.italic = True
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    sub_run.font.size = Pt(10)

    for buyer in sorted_buyers:
        items = by_buyer[buyer]
        recs = sorted(recipients.get(buyer, set()))

        # 구매자 헤더
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(buyer)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)

        # 수취인이 다르면 작게 표시
        if not (len(recs) == 1 and recs[0] == buyer):
            run2 = p.add_run(f"   (수취인: {', '.join(recs)})")
            run2.font.size = Pt(10)
            run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # 표
        table = doc.add_table(rows=1 + len(items), cols=2)
        table.autofit = False
        widths = (Cm(13.5), Cm(2.5))
        # 표 너비 강제
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = widths[i]
                set_cell_borders(cell)

        # 헤더
        hdr = table.rows[0]
        hdr.cells[0].text = ""
        hdr.cells[1].text = ""
        for i, txt in enumerate(["Product", "Qty"]):
            cell = hdr.cells[i]
            shade_cell(cell, "EAEAEA")
            para = cell.paragraphs[0]
            run = para.add_run(txt)
            run.bold = True
            if i == 1:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 데이터 행
        for ri, it in enumerate(items, 1):
            row = table.rows[ri]
            row.cells[0].text = ""
            row.cells[1].text = ""
            row.cells[0].paragraphs[0].add_run(it["eng"])
            qpara = row.cells[1].paragraphs[0]
            qpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
            qpara.add_run(str(it["qty"]))
            row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.save(out_path)
    print(f"DONE:{out_path}:{len(sorted_buyers)}:{sum(len(v) for v in by_buyer.values())}")
    if missing_set:
        uniq = list(dict.fromkeys(missing_set))  # preserve order, unique
        print("MISSING_NAMES:" + "|".join(uniq))


if __name__ == "__main__":
    main()
